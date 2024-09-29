# chapter_generator.py
import os
from typing import Dict, Optional, Any, List, Tuple, Union
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document  # type: ignore
import logging
import traceback
import datetime

# Load environment variables from .env file
load_dotenv()

class ChapterGenerator:
    def __init__(self, api_key: str, generation_model: str, check_model: str):
        genai.configure(api_key=api_key)
        self.generation_model = genai.GenerativeModel(generation_model)
        self.check_model = genai.GenerativeModel(check_model)
        self.MAX_INPUT_TOKENS = 2097152 if 'pro' in generation_model else 1048576
        self.MAX_OUTPUT_TOKENS = 8192
        self.logger = logging.getLogger(__name__)

    def generate_chapter(self, instructions: Dict[str, Any], context: str, chapter_path: str, chapter_number: int) -> str:
        try:
            self.logger.info(f"Generating chapter {chapter_number}...")

            previous_chapters = self.get_existing_chapter_content(chapter_number)
            
            prompt = self._construct_prompt(instructions, context, previous_chapters, chapter_path)
            
            self.logger.debug(f"Generated prompt for chapter {chapter_number}")
            
            generation_config = genai.GenerationConfig(
                max_output_tokens=self.MAX_OUTPUT_TOKENS,
                temperature=1,
            )
            
            response = self.generation_model.generate_content(
                prompt,
                generation_config=generation_config,
                stream=True
            )   
            
            chapter = ""
            for chunk in response:
                chapter += chunk.text
            
            self.logger.info(f"Chapter {chapter_number} generation completed.")
            
            # Check the chapter
            self.logger.info(f"Checking chapter {chapter_number}...")
            is_valid, feedback = self.check_chapter(chapter, instructions, previous_chapters)
            
            # Review the chapter
            review_feedback = self.review_chapter(chapter, instructions, previous_chapters)

            # Enforce style guide
            style_guide = instructions.get('style_guide', '')
            adheres_to_style_guide, style_guide_feedback = self.enforce_style_guide(chapter, style_guide)

            # Check continuity (only if there are previous chapters)
            continuity = True
            continuity_feedback = "No previous chapters available for continuity check."
            if previous_chapters:
                continuity, continuity_feedback = self.check_continuity(chapter, previous_chapters)

            # Run automated tests
            test_results = self.run_tests(chapter, instructions, previous_chapters)

            # Save the chapter regardless of warnings
            self.save_response(chapter, chapter_path, chapter_number)
            self.save_validity_feedback(is_valid, feedback, chapter_path, review_feedback, style_guide_feedback, continuity_feedback, test_results, adheres_to_style_guide, continuity, chapter_number)

            return chapter

        except Exception as e:
            self.logger.error(f"Error generating chapter {chapter_number}: {e}")
            self.logger.error(traceback.format_exc())
            return f"Error generating chapter: {str(e)}"

    def _get_chapter_number(self, chapter_path: str) -> int:
        try:
            chapter_number = int(chapter_path.split('/')[-1].split(' ')[1].split('.')[0])
            return chapter_number
        except (ValueError, IndexError):
            self.logger.warning(f"Could not parse chapter number from file path: {chapter_path}. Defaulting to 1.")
            return 1  # Default to chapter 1 if extraction fails

    def _construct_prompt(self, instructions: Dict[str, Any], context: str, previous_chapters: Optional[str], chapter_path: str) -> str:
        prompt = f"""
        Instructions:
        Plot: {instructions.get('plot', '')}
        Writing Style: {instructions.get('writing_style', '')}
        Additional Instructions: {instructions.get('instructions', '')}
        Chapter Filename: {chapter_path}
        
        Context: {context}
        
        Previous Chapters: {previous_chapters if previous_chapters else "No previous chapters"}...
        
        Based on the above instructions, context, and previous chapters, write a new chapter for the story. 
        Ensure that the chapter follows the plot points, incorporates the characters and settings, 
        adheres to the specified writing style, and maintains continuity with previous chapters.
        """
        
        return prompt

    def check_chapter(self, chapter: str, instructions: Dict[str, Any], previous_chapters: Optional[str]) -> Tuple[bool, str]:
        try:
            prompt = f"""
            Chapter: {chapter}
            Instructions:
            Plot: {instructions.get('plot', '')}
            Writing Style: {instructions.get('writing_style', '')}
            Additional Instructions: {instructions.get('instructions', '')}
            Chapter Filename: {instructions.get('chapter_filename', '')}  # Add chapter filename to the prompt
            
            Previous Chapters: {previous_chapters}... 
            
            Check if the chapter is valid based on the instructions, previous chapters, and context embedding. 
            Provide feedback on adherence to plot, character development, setting descriptions, 
            and writing style. Return your response in the following format:
            
            Valid: [Yes/No]
            Feedback: [Your detailed feedback here]
            """
            response = self.check_model.generate_content(prompt)
            response_text = response.text
            is_valid = "Valid: Yes" in response_text
            feedback = response_text.split("Feedback:")[1].strip() if "Feedback:" in response_text else response_text
            return is_valid, feedback
        except Exception as e:
            self.logger.error(f"Error checking chapter: {e}")
            return False, "An error occurred while checking the chapter."

    def review_chapter(self, chapter: str, instructions: Dict[str, Any], previous_chapters: Optional[str]) -> List[str]:
        try:
            prompt = f"""
            Chapter: {chapter}
            Instructions:
            Plot: {instructions.get('plot', '')}
            Writing Style: {instructions.get('writing_style', '')}
            Additional Instructions: {instructions.get('instructions', '')}
            Chapter Filename: {instructions.get('chapter_filename', '')}  # Add chapter filename to the prompt
            
            Previous Chapters: {previous_chapters}
            
            Provide a detailed review of the chapter, focusing on plot consistency, character development, setting descriptions, 
            and adherence to the specified writing style. Return your response in the following format:
            
            Review: [Your detailed review here]
            """
            response = self.check_model.generate_content(prompt)
            response_text = response.text
            review_feedback = response_text.split("Review:")[1].strip() if "Review:" in response_text else response_text
            return review_feedback.split('\n')
        except Exception as e:
            self.logger.error(f"Error reviewing chapter: {e}")
            return ["An error occurred while reviewing the chapter."]

    def enforce_style_guide(self, chapter: str, style_guide: str) -> Tuple[bool, str]:
        try:
            prompt = f"""
            Chapter: {chapter}
            Style Guide: {style_guide}
            
            Check if the chapter adheres to the specified style guide. Provide feedback on any deviations. 
            Return your response in the following format:
            
            Adheres to Style Guide: [Yes/No]
            Feedback: [Your detailed feedback here]
            """
            response = self.check_model.generate_content(prompt)
            response_text = response.text
            adheres_to_style_guide = "Adheres to Style Guide: Yes" in response_text
            feedback = response_text.split("Feedback:")[1].strip() if "Feedback:" in response_text else response_text
            return adheres_to_style_guide, feedback
        except Exception as e:
            self.logger.error(f"Error enforcing style guide: {e}")
            return False, "An error occurred while enforcing the style guide."

    def check_continuity(self, chapter: str, previous_chapters: str) -> Tuple[bool, str]:
        try:
            prompt = f"""
            New Chapter: {chapter}
            Previous Chapters: {previous_chapters}
            
            Check if the new chapter maintains continuity with the previous chapters. 
            Provide feedback on any inconsistencies. Return your response in the following format:
            
            Continuity: [Yes/No]
            Feedback: [Your detailed feedback here]
            """
            response = self.check_model.generate_content(prompt)
            response_text = response.text
            continuity = "Continuity: Yes" in response_text
            feedback = response_text.split("Feedback:")[1].strip() if "Feedback:" in response_text else response_text
            return continuity, feedback
        except Exception as e:
            self.logger.error(f"Error checking continuity: {e}")
            return False, "An error occurred while checking continuity."

    def get_existing_chapter_content(self, chapter_number: int) -> Optional[str]:
        existing_content = ""
        total_tokens = 0
        
        for i in range(1, chapter_number):
            previous_chapter_path = f'output/Chapter {i}.docx'
            if os.path.exists(previous_chapter_path):
                doc = Document(previous_chapter_path)
                chapter_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                chapter_tokens = self.estimate_token_count(chapter_content)
                
                if total_tokens + chapter_tokens > self.MAX_INPUT_TOKENS // 2:  # Use only half for previous chapters
                    break
                
                existing_content = chapter_content + "\n" + existing_content
                total_tokens += chapter_tokens
        
        return existing_content if existing_content else None

    def save_response(self, chapter: str, output_path: str, chapter_number: int):
        # Increment chapter number if file already exists
        while os.path.exists(output_path):
            chapter_number += 1
            output_path = f'output/Chapter {chapter_number}.docx'
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = Document()
        doc.add_heading(f'Chapter {chapter_number}', level=1)
        doc.add_paragraph(chapter)
        doc.save(output_path)

    def save_validity_feedback(self, is_valid: bool, feedback: str, chapter_path: str, review_feedback: List[str], style_guide_feedback: str, continuity_feedback: str, test_results: List[str], adheres_to_style_guide: bool, continuity: bool, chapter_number: int):
        validity_path = f'output/Chapter {chapter_number} validity.docx'
        os.makedirs(os.path.dirname(validity_path), exist_ok=True)
        doc = Document()
        doc.add_heading(f'Chapter Validity and Feedback', level=1)
        doc.add_paragraph(f"Is Valid: {is_valid}")
        doc.add_paragraph(f"Feedback: {feedback}")
        if review_feedback:
            doc.add_heading(f'Content Review Feedback', level=2)
            for item in review_feedback:
                doc.add_paragraph(item)
        if style_guide_feedback:
            doc.add_heading(f'Style Guide Feedback', level=2)
            doc.add_paragraph(f"Adheres to Style Guide: {adheres_to_style_guide}")
            doc.add_paragraph(f"Feedback: {style_guide_feedback}")
        if continuity_feedback:
            doc.add_heading(f'Continuity Feedback', level=2)
            doc.add_paragraph(f"Continuity: {continuity}")
            doc.add_paragraph(f"Feedback: {continuity_feedback}")
        if test_results:
            doc.add_heading(f'Automated Test Results', level=2)
            for result in test_results:
                doc.add_paragraph(result)
        doc.save(validity_path)

    def run_tests(self, chapter: str, instructions: Dict[str, Any], previous_chapters: Optional[str]) -> List[str]:
        try:
            prompt = f"""
            Chapter: {chapter}
            Instructions:
            Plot: {instructions.get('plot', '')}
            Writing Style: {instructions.get('writing_style', '')}
            Additional Instructions: {instructions.get('instructions', '')}
            Chapter Filename: {instructions.get('chapter_filename', '')}  # Add chapter filename to the prompt
            
            Previous Chapters: {previous_chapters}
            
            Run automated tests on the chapter to ensure it meets the specified criteria. 
            Provide feedback on any issues found. Return your response in the following format:
            
            Test Results: [Your detailed test results here]
            """
            response = self.check_model.generate_content(prompt)
            response_text = response.text
            test_results = response_text.split("Test Results:")[1].strip() if "Test Results:" in response_text else response_text
            return test_results.split('\n')
        except Exception as e:
            self.logger.error(f"Error running tests: {e}")
            return ["An error occurred while running tests."]

    def estimate_token_count(self, text: str) -> int:
        # Gemini models use about 4 characters per token
        return len(text) // 4

    def embed_content(self, content: str) -> List[float]:
        result = genai.embed_content(
            model="models/text-embedding-004",
            content=content,
            task_type="retrieval_document",
            title="Chapter embedding"
        )
        return result['embedding']

    def get_relevant_context(self, query: str, previous_chapters: Optional[List[str]]) -> str:
        query_embedding = self.embed_content(query)
        
        if previous_chapters:
            chapter_embeddings = [self.embed_content(chapter) for chapter in previous_chapters]
            
            # Implement similarity search here (e.g., cosine similarity)
            # Return the most relevant previous chapter content
            
            # For simplicity, just returning the last chapter here
            return previous_chapters[-1]
        else:
            return ""
